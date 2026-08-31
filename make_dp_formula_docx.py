from html import escape
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = Path(r"C:\Users\31491\Desktop\2\状态转移与代价函数公式整理.docx")

MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, size=11, color=None, bold=False):
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor.from_string(color)


def patch_paragraph_spacing(style, before_pt=0, after_pt=6, line=300):
    ppr = style._element.get_or_add_pPr()
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    spacing.set(qn("w:before"), str(int(before_pt * 20)))
    spacing.set(qn("w:after"), str(int(after_pt * 20)))
    spacing.set(qn("w:line"), str(line))
    spacing.set(qn("w:lineRule"), "auto")


def mr(text, italic=False):
    text = escape(text)
    style = '<m:rPr><m:sty m:val="i"/></m:rPr>' if italic else ""
    return f'<m:r>{style}<m:t xml:space="preserve">{text}</m:t></m:r>'


def var(text):
    return mr(text, italic=True)


def op(text):
    return mr(text, italic=False)


def sub(base_xml, sub_xml):
    return f"<m:sSub><m:e>{base_xml}</m:e><m:sub>{sub_xml}</m:sub></m:sSub>"


def sup(base_xml, sup_xml):
    return f"<m:sSup><m:e>{base_xml}</m:e><m:sup>{sup_xml}</m:sup></m:sSup>"


def frac(num_xml, den_xml):
    return (
        '<m:f><m:fPr><m:type m:val="bar"/></m:fPr>'
        f"<m:num>{num_xml}</m:num><m:den>{den_xml}</m:den></m:f>"
    )


def math_el(inner_xml):
    return parse_xml(f'<m:oMath xmlns:m="{MATH_NS}">{inner_xml}</m:oMath>')


def add_math(paragraph, inner_xml):
    paragraph._p.append(math_el(inner_xml))


def add_math_line(doc, inner_xml):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph._p.append(math_el(inner_xml))
    return paragraph


def add_bullet(doc, math_xml=None, chunks=None):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    if math_xml:
        add_math(paragraph, math_xml)
    for chunk in chunks or []:
        if isinstance(chunk, tuple):
            add_math(paragraph, chunk[1])
        else:
            run = paragraph.add_run(chunk)
            set_run_font(run, size=11)
    return paragraph


def eq_soc_range():
    return (
        var("SOC")
        + op(" ∈ [")
        + sub(var("SOC"), var("min"))
        + op(", ")
        + sub(var("SOC"), var("max"))
        + op("]")
    )


def eq_temp_range():
    return (
        var("T")
        + op(" ∈ [")
        + sub(var("T"), var("low"))
        + op(", ")
        + sub(var("T"), var("high"))
        + op(" + Δ]")
    )


def eq_peng_range():
    return (
        sub(var("P"), var("eng"))
        + op(" ∈ [0, ")
        + sub(var("P"), var("eng_max"))
        + op("]")
    )


def eq_pcool():
    return sub(var("P"), var("cool"))


def eq_t_high_switch():
    return var("T") + op(" ≥ ") + sub(var("T"), var("high"))


def eq_pbat():
    return (
        sub(var("P"), var("bat"))
        + op(" = ")
        + sub(var("P"), var("req_total"))
        + op("(t) − ")
        + sub(var("P"), var("eng"))
    )


def eq_soc_transition():
    numerator = sub(var("I"), var("bat")) + op(" · dt")
    denominator = op("3600 · ") + var("Cap")
    return (
        sub(var("SOC"), op("k+1"))
        + op(" = ")
        + sub(var("SOC"), var("k"))
        + op(" − ")
        + frac(numerator, denominator)
    )


def eq_temp_transition():
    ibat_sq = sup(sub(var("I"), var("bat")), op("2"))
    numerator = (
        ibat_sq
        + sub(var("R"), op("0"))
        + op(" − ")
        + sub(var("P"), var("cool"))
        + op(" · 1000")
    )
    denominator = sub(var("C"), var("th"))
    return (
        sub(var("T"), op("k+1"))
        + op(" = ")
        + sub(var("T"), var("k"))
        + op(" + dt · ")
        + frac(numerator, denominator)
    )


def eq_cost():
    return (
        sub(var("Cost"), var("k"))
        + op(" = ")
        + var("Fuel")
        + op("(")
        + sub(var("P"), var("eng"))
        + op(") + ")
        + sub(var("w"), var("elec"))
        + op(" · ")
        + sub(var("P"), var("bat"))
        + op(" + ")
        + sub(var("w"), var("cool"))
        + op(" · ")
        + sub(var("P"), var("cool"))
    )


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
set_style_font(styles["Normal"], size=11)
patch_paragraph_spacing(styles["Normal"], before_pt=0, after_pt=6, line=300)
set_style_font(styles["Heading 1"], size=16, color="2E74B5", bold=True)
patch_paragraph_spacing(styles["Heading 1"], before_pt=18, after_pt=10, line=300)
set_style_font(styles["Heading 2"], size=13, color="2E74B5", bold=True)
patch_paragraph_spacing(styles["Heading 2"], before_pt=14, after_pt=7, line=300)
set_style_font(styles["List Bullet"], size=11)
patch_paragraph_spacing(styles["List Bullet"], before_pt=0, after_pt=4, line=300)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(12)
run = title.add_run("动态规划状态转移与代价函数公式整理")
set_run_font(run, size=16, bold=True, color="1F4D78")

doc.add_heading("1. 状态变量离散化", level=2)
add_bullet(doc, eq_soc_range(), ["，划分为 ", ("math", sub(var("N"), var("s"))), " 个网格点。"])
add_bullet(doc, eq_temp_range(), ["，划分为 ", ("math", sub(var("N"), var("T"))), " 个网格点。"])

doc.add_heading("2. 控制变量离散化", level=2)
add_bullet(doc, eq_peng_range(), ["，划分为 ", ("math", sub(var("N"), var("p"))), " 个网格点。"])

doc.add_heading("3. 状态转移方程", level=2)
add_bullet(
    doc,
    eq_pcool(),
    [
        " 根据当前 ",
        ("math", var("T")),
        " 决定（滞回控制简化为：",
        ("math", eq_t_high_switch()),
        " 时满功率，否则为 0）。",
    ],
)
add_bullet(doc, eq_pbat())
add_bullet(doc, eq_soc_transition())
add_bullet(doc, eq_temp_transition())

doc.add_heading("4. 代价函数", level=2)
add_bullet(doc, eq_cost())

doc.save(OUT_PATH)
print(OUT_PATH)
